import streamlit as st
import requests
from docx import Document
from io import BytesIO

# Función para corregir texto usando LanguageTool API
def corregir_texto(texto, idioma='es'):
    url = "https://api.languagetool.org/v2/check"
    payload = {
        'text': texto,
        'language': idioma,
        'enabledOnly': False,
    }
    response = requests.post(url, data=payload)
    if response.status_code != 200:
        st.error("Error al conectar con la API de LanguageTool.")
        return texto  # Retorna el texto original si hay error

    matches = response.json().get('matches', [])
    correcciones = []

    for match in matches:
        offset = match['offset']
        length = match['length']
        replacement = match['replacements'][0]['value'] if match['replacements'] else ''
        correcciones.append((offset, length, replacement))

    # Aplicar las correcciones de atrás hacia adelante para no afectar los offsets
    correcciones = sorted(correcciones, key=lambda x: x[0], reverse=True)
    texto_modificado = texto
    for offset, length, replacement in correcciones:
        texto_modificado = texto_modificado[:offset] + replacement + texto_modificado[offset + length:]

    return texto_modificado

# Función para procesar el documento de Word
def procesar_documento(doc_bytes):
    doc = Document(BytesIO(doc_bytes))
    doc_corregido = Document()

    for para in doc.paragraphs:
        texto_original = para.text
        if texto_original.strip() != '':
            texto_corregido = corregir_texto(texto_original)
        else:
            texto_corregido = texto_original

        # Copiar el estilo del párrafo original
        nuevo_para = doc_corregido.add_paragraph()
        nuevo_para.style = para.style
        nuevo_para.add_run(texto_corregido)

    # Copiar el contenido de las tablas si las hay
    for table in doc.tables:
        nueva_tabla = doc_corregido.add_table(rows=0, cols=0)
        for row in table.rows:
            nueva_fila = nueva_tabla.add_row()
            for idx, cell in enumerate(row.cells):
                if idx >= len(nueva_fila.cells):
                    nueva_fila.add_cell()
                texto_original = cell.text
                if texto_original.strip() != '':
                    texto_corregido = corregir_texto(texto_original)
                else:
                    texto_corregido = texto_original
                nueva_fila.cells[idx].text = texto_corregido

    # Guardar el documento corregido en un buffer
    buffer = BytesIO()
    doc_corregido.save(buffer)
    buffer.seek(0)
    return buffer

# Configuración de la aplicación Streamlit
st.title("Corrector de Ortografía y Gramática para Documentos Word")
st.write("""
    Sube un archivo de Word (.docx) y la aplicación corregirá la ortografía y gramática utilizando la API de LanguageTool.
    Luego, podrás descargar el documento corregido sin modificar el formato original.
""")

# Carga de archivo
uploaded_file = st.file_uploader("Selecciona un archivo de Word (.docx)", type=["docx"])

if uploaded_file is not None:
    st.success("Archivo cargado exitosamente.")
    if st.button("Corregir Documento"):
        with st.spinner("Corrigiendo el documento..."):
            try:
                buffer_corregido = procesar_documento(uploaded_file.read())
                st.success("Documento corregido exitosamente.")
                st.download_button(
                    label="Descargar Documento Corregido",
                    data=buffer_corregido,
                    file_name="documento_corregido.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"Ocurrió un error al procesar el documento: {e}")
